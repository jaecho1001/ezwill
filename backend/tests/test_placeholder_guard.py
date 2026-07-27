"""Unresolved-placeholder guard tests.

The worst failure mode for a legal-document product is silent corruption: a
will delivered with literal [executorName] text where the client's data should
be. These tests pin the guard at every layer — the resolve_variables collector,
the assembled-document scan for the bracket literals that bypass it, and the
422 rejection (with explicit override) on the delivery routes.
"""

from __future__ import annotations

import io
import os
import sys
import zipfile

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

# Module-level import on purpose: test_routes.py replaces sys.modules['docx']
# with a MagicMock during collection, so a runtime import inside a test body
# would get the mock. Binding the real Document here (this module imports
# before test_routes alphabetically) keeps these tests on real python-docx.
from docx import Document as RealDocument

from services.document_generator import (
    DocumentGenerator,
    resolve_variables,
    scan_unresolved_literals,
)
from routes import documents as documents_route
from routes.auth import verify_dashboard_token


# ── resolve_variables collector ──────────────────────────────────────────────

def test_missing_placeholder_is_recorded():
    missing: set = set()
    out = resolve_variables("I appoint {{executorName}}.", {}, missing)
    assert out == "I appoint [executorName]."
    assert missing == {"executorName"}


def test_camel_to_snake_fallback_is_not_recorded_as_missing():
    missing: set = set()
    out = resolve_variables(
        "{{testatorFullName}}", {"testator_full_name": "HYUN JUNG KIM"}, missing
    )
    assert out == "HYUN JUNG KIM"
    assert missing == set()


def test_empty_string_value_is_recorded_as_missing():
    # A present-but-blank variable substitutes to "" (unchanged) but IS
    # recorded: "my spouse, ." in a signed will is the exact silent
    # corruption the guard exists to refuse — worse than a bracket, since
    # there is no visible marker at all. (_build_variables defaults many
    # keys to "", which previously made this blind spot structural.)
    missing: set = set()
    assert resolve_variables("{{spouseFullName}}", {"spouseFullName": ""}, missing) == ""
    assert missing == {"spouseFullName"}


def test_whitespace_only_value_is_recorded_as_missing():
    missing: set = set()
    assert resolve_variables("{{city}}", {"city": "  "}, missing) == "  "
    assert missing == {"city"}


def test_blank_value_without_collector_still_substitutes():
    # Preview/review-portal callers without a collector see no behavior change.
    assert resolve_variables("{{spouseFullName}}", {"spouseFullName": ""}) == ""


def test_none_text_passthrough_preserved():
    missing: set = set()
    assert resolve_variables(None, {}, missing) is None
    assert missing == set()


def test_collector_is_optional():
    # Preview/review-portal callers pass no collector; behavior is unchanged.
    assert resolve_variables("{{x}}", {}) == "[x]"


# ── assembled-document scan (cover/signing pages bypass resolve_variables) ──

def _doc_text(docx_bytes: bytes) -> str:
    doc = RealDocument(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


CLAUSE_MISSING = [{
    "clause_id": "executors",
    "included": True,
    "templateText": "I appoint {{primaryExecutorFullName}} as my Estate Trustee.",
    "sortOrder": 1,
    "title": "Appointment of Estate Trustee",
}]

CLAUSE_COMPLETE = [{
    "clause_id": "revocation",
    "included": True,
    "templateText": "I, {{testatorFullName}}, revoke all prior wills.",
    "sortOrder": 1,
    "title": "Revocation",
}]

RESIDUE_CLAUSE_MATCHING = [{
    "clause_id": "res-named-shares",
    "included": True,
    "customText": (
        "I give 60% of the residue of my estate to Alex Kim and "
        "40% to Sam Lee."
    ),
    "section": "Residue",
    "sortOrder": 1,
    "title": "Named residue gifts",
}]

POA_CARE_WISHES = [{
    "clause_id": "poa-care-wishes",
    "included": True,
    "templateText": "My attorney shall withhold life-sustaining treatment.",
    "sortOrder": 1,
    "title": "End-of-life wishes",
}]

FULL_VARIABLES = {
    "testatorFullName": "HYUN JUNG KIM",
    "documentDate": "July 19, 2026",
}


def test_generate_document_collects_clause_misses():
    missing: set = set()
    DocumentGenerator().generate_document(
        "single_will", CLAUSE_MISSING, dict(FULL_VARIABLES), missing=missing
    )
    assert "primaryExecutorFullName" in missing


def test_signing_page_fallback_literal_is_caught():
    # No testator name anywhere: the signing/cover pages emit their hardcoded
    # bracket literals without ever calling resolve_variables. The final-doc
    # scan must still catch them.
    missing: set = set()
    docx_bytes = DocumentGenerator().generate_document(
        "single_will", CLAUSE_COMPLETE, {}, missing=missing
    )
    text = _doc_text(docx_bytes)
    assert "[TESTATOR NAME]" in text or "[Client Name]" in text
    assert "[TESTATOR NAME]" in missing or "[Client Name]" in missing


def test_unmatchable_brace_token_is_caught():
    # {{ spaced token }} never matches VARIABLE_PATTERN, survives substitution
    # verbatim, and must be reported by the scan.
    clause = [{
        "clause_id": "bad",
        "included": True,
        "templateText": "Residue to {{ residual beneficiary }}.",
        "sortOrder": 1,
        "title": "Residue",
    }]
    missing: set = set()
    DocumentGenerator().generate_document(
        "single_will", clause, dict(FULL_VARIABLES), missing=missing
    )
    assert "{{ residual beneficiary }}" in missing


def test_complete_document_reports_nothing():
    missing: set = set()
    DocumentGenerator().generate_document(
        "single_will", CLAUSE_COMPLETE, dict(FULL_VARIABLES), missing=missing
    )
    assert missing == set()


def test_scan_helper_finds_literals_in_tables():
    doc = RealDocument()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].add_run("Witness: [DEPONENT NAME]")
    missing: set = set()
    scan_unresolved_literals(doc, missing)
    assert missing == {"[DEPONENT NAME]"}


def test_commissioner_blank_is_fill_at_signing_not_flagged():
    # No code path can populate a commissioner at generation time — that
    # block is completed at commissioning. Flagging it would 422 every
    # affidavit ever generated.
    doc = RealDocument()
    doc.add_paragraph("A Commissioner for Taking Affidavits: [COMMISSIONER NAME]")
    missing: set = set()
    scan_unresolved_literals(doc, missing)
    assert missing == set()


def test_affidavit_with_witness_and_city_passes_guard():
    variables = {
        "testatorFullName": "HYUN JUNG KIM",
        "otherWitnessName": "JOON TAE PARK",
        "city": "Vaughan",
        "documentDate": "July 19, 2026",
        "province": "ON",
        "numberOfPages": "12",
    }
    missing: set = set()
    DocumentGenerator().generate_document(
        "affidavit_execution", CLAUSE_COMPLETE, variables, missing=missing
    )
    assert missing == set()


def test_affidavit_without_deponent_is_flagged():
    variables = {
        "testatorFullName": "HYUN JUNG KIM",
        "city": "Vaughan",
        "documentDate": "July 19, 2026",
        "province": "ON",
        "numberOfPages": "12",
    }
    missing: set = set()
    DocumentGenerator().generate_document(
        "affidavit_execution", CLAUSE_COMPLETE, variables, missing=missing
    )
    assert "[DEPONENT NAME]" in missing


# ── route-level rejection ────────────────────────────────────────────────────

class FakeDb:
    """Stands in for EWDbWriter inside routes.documents."""

    generated_records: list = []
    clause_selections: list = []
    all_selections: dict = {}

    def __init__(self, schema: str):
        self.schema = schema

    def __enter__(self):
        return self

    def __exit__(self, *_exc):
        return None

    def get_clause_selections(self, draft_id, document_type):
        return list(type(self).clause_selections)

    def get_all_clause_selections(self, draft_id):
        return dict(type(self).all_selections)

    def get_document_configs(self, draft_id):
        return []

    def get_firm_settings(self):
        return {}

    def record_document_generation(self, draft_id, document_type, file_format,
                                   content, generated_by="dashboard", params=None):
        return {
            "id": "gen-1", "created_at": "now",
            "storage_path": "db://ew_document_generations/gen-1",
            "content_sha256": "x", "byte_size": len(content),
        }

    def update_document_generated(self, draft_id, document_type, file_path):
        type(self).generated_records.append((draft_id, document_type, file_path))
        return True


DRAFT = {
    "id": "draft-1",
    "client_first_name": "Hyun Jung",
    "client_last_name": "Kim",
    "province": "ON",
    "people": [],
}


@pytest.fixture
def client(monkeypatch):
    monkeypatch.setattr(documents_route, "EWDbWriter", FakeDb)
    monkeypatch.setattr(documents_route, "get_full_draft", lambda draft_id, schema: dict(DRAFT))
    FakeDb.generated_records = []

    app = FastAPI()
    app.include_router(documents_route.router, prefix="/api/documents")
    app.dependency_overrides[verify_dashboard_token] = lambda: "test-token"
    return TestClient(app)


def test_generate_rejects_unresolved_placeholders(client):
    FakeDb.clause_selections = CLAUSE_MISSING
    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx"},
    )
    assert res.status_code == 422
    detail = res.json()["detail"]
    assert detail["error"] == "unresolved_placeholders"
    assert "primaryExecutorFullName" in detail["unresolved"]
    # Nothing recorded as generated: the document was never delivered.
    assert FakeDb.generated_records == []


def test_generate_override_delivers_with_warning(client):
    FakeDb.clause_selections = CLAUSE_MISSING
    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx", "allow_incomplete": True},
    )
    assert res.status_code == 200
    assert res.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert len(FakeDb.generated_records) == 1


def test_generate_complete_document_passes(client):
    FakeDb.clause_selections = CLAUSE_COMPLETE
    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx"},
    )
    assert res.status_code == 200
    assert len(FakeDb.generated_records) == 1


def test_generate_rejects_residue_clause_that_omits_client_shares(
    client, monkeypatch
):
    FakeDb.clause_selections = CLAUSE_COMPLETE
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={
                "beneficiaries": [
                    {"fullName": "Alex Kim", "sharePercent": 60},
                    {"fullName": "Sam Lee", "sharePercent": 40},
                ],
                "residueDistribution": "percentages",
            },
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx"},
    )

    assert res.status_code == 422
    detail = res.json()["detail"]
    assert (
        documents_route.RESIDUE_INSTRUCTION_GAP
        in detail["instruction_gaps"]
    )
    assert FakeDb.generated_records == []


def test_generate_accepts_custom_residue_clause_matching_client_shares(
    client, monkeypatch
):
    FakeDb.clause_selections = RESIDUE_CLAUSE_MATCHING
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={
                "beneficiaries": [
                    {"fullName": "Alex Kim", "sharePercent": 60},
                    {"fullName": "Sam Lee", "sharePercent": 40},
                ],
                "residueDistribution": "percentages",
            },
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx"},
    )

    assert res.status_code == 200
    assert len(FakeDb.generated_records) == 1


def test_generate_rejects_legacy_residue_instructions_without_matching_clause(
    client, monkeypatch
):
    FakeDb.clause_selections = CLAUSE_COMPLETE
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={},
            your_estate={
                "beneficiaries": [
                    {"firstName": "Alex", "lastName": "Kim", "percentage": 60},
                    {"firstName": "Sam", "lastName": "Lee", "percentage": 40},
                ],
                "residueDistribution": "custom",
            },
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx"},
    )

    assert res.status_code == 422
    assert (
        documents_route.RESIDUE_INSTRUCTION_GAP
        in res.json()["detail"]["instruction_gaps"]
    )


def test_generate_uses_stored_legacy_beneficiary_rows_as_guard_fallback(
    client, monkeypatch
):
    FakeDb.clause_selections = CLAUSE_COMPLETE
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={},
            your_estate={"residueDistribution": "custom"},
            people=[
                {
                    "role": "beneficiary",
                    "first_name": "Alex",
                    "last_name": "Kim",
                    "percentage": 100,
                },
            ],
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx"},
    )

    assert res.status_code == 422
    assert (
        documents_route.RESIDUE_INSTRUCTION_GAP
        in res.json()["detail"]["instruction_gaps"]
    )


def test_generate_rejects_residue_clause_with_swapped_percentages(
    client, monkeypatch
):
    FakeDb.clause_selections = [{
        **RESIDUE_CLAUSE_MATCHING[0],
        "customText": (
            "I give 40% of the residue of my estate to Alex Kim and "
            "60% to Sam Lee."
        ),
    }]
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={
                "beneficiaries": [
                    {"fullName": "Alex Kim", "sharePercent": 60},
                    {"fullName": "Sam Lee", "sharePercent": 40},
                ],
                "residueDistribution": "percentages",
            },
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx"},
    )

    assert res.status_code == 422
    assert (
        documents_route.RESIDUE_INSTRUCTION_GAP
        in res.json()["detail"]["instruction_gaps"]
    )


def test_child_per_stirpes_choice_ignores_stale_named_beneficiary_list(
    client, monkeypatch
):
    FakeDb.clause_selections = CLAUSE_COMPLETE
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={
                "beneficiaries": [
                    {"fullName": "Old beneficiary", "sharePercent": 100},
                ],
                "residueDistribution": "per_stirpes",
            },
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "single_will", "format": "docx"},
    )

    assert res.status_code == 200


def test_generate_rejects_end_of_life_clause_without_client_instruction(
    client, monkeypatch
):
    FakeDb.clause_selections = POA_CARE_WISHES
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={"poa": {"personalCare": {"lifeSupport": "unsure"}}},
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "poa_personal_care", "format": "docx"},
    )

    assert res.status_code == 422
    detail = res.json()["detail"]
    assert (
        documents_route.END_OF_LIFE_INSTRUCTION_GAP
        in detail["instruction_gaps"]
    )
    assert FakeDb.generated_records == []


def test_generate_accepts_end_of_life_clause_after_affirmative_instruction(
    client, monkeypatch
):
    FakeDb.clause_selections = POA_CARE_WISHES
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={"poa": {"personalCare": {"lifeSupport": "withhold"}}},
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "poa_personal_care", "format": "docx"},
    )

    assert res.status_code == 200
    assert len(FakeDb.generated_records) == 1


def test_generate_accepts_legacy_end_of_life_instruction(
    client, monkeypatch
):
    FakeDb.clause_selections = POA_CARE_WISHES
    monkeypatch.setattr(
        documents_route,
        "get_full_draft",
        lambda draft_id, schema: dict(
            DRAFT,
            vault={},
            poa_personal_care={"lifeSupport": "withhold"},
        ),
    )

    res = client.post(
        "/api/documents/draft-1/generate",
        json={"document_type": "poa_personal_care", "format": "docx"},
    )

    assert res.status_code == 200


def test_generate_all_rejects_batch_when_any_document_incomplete(client):
    FakeDb.all_selections = {
        "single_will": CLAUSE_COMPLETE,
        "poa_property": CLAUSE_MISSING,
    }
    res = client.post("/api/documents/draft-1/generate-all")
    assert res.status_code == 422
    detail = res.json()["detail"]
    assert "poa_property" in detail["unresolved_by_document"]
    assert "single_will" not in detail["unresolved_by_document"]
    assert FakeDb.generated_records == []


def test_generate_all_override_delivers_zip(client):
    FakeDb.all_selections = {
        "single_will": CLAUSE_COMPLETE,
        "poa_property": CLAUSE_MISSING,
    }
    res = client.post("/api/documents/draft-1/generate-all?allow_incomplete=true")
    assert res.status_code == 200
    with zipfile.ZipFile(io.BytesIO(res.content)) as zf:
        assert len(zf.namelist()) == 2


def test_preview_reports_unresolved_but_does_not_block(client):
    FakeDb.clause_selections = CLAUSE_MISSING
    res = client.get("/api/documents/draft-1/preview/single_will")
    assert res.status_code == 200
    body = res.json()
    assert "primaryExecutorFullName" in body["unresolved_placeholders"]


def test_preview_warning_list_matches_generation(client, monkeypatch):
    # Parity both directions (review finding): a folder row's unused
    # templateText must NOT be reported, while the cover/signing-page bracket
    # literals — which the HTML preview never renders — MUST be, so the
    # preview's warning list predicts exactly what POST /generate refuses.
    FakeDb.clause_selections = [
        {
            "clause_id": "folder-1", "included": True, "isFolder": True,
            "title": "EXECUTORS", "templateText": "{{neverRenderedInFolders}}",
            "sortOrder": 1,
        },
        dict(CLAUSE_COMPLETE[0], sortOrder=2),
    ]
    # Draft with no client name: generation emits [Client Name]/[TESTATOR NAME]
    # fallback literals on cover/signing pages.
    monkeypatch.setattr(
        documents_route, "get_full_draft",
        lambda draft_id, schema: dict(DRAFT, client_first_name="", client_last_name=""),
    )
    res = client.get("/api/documents/draft-1/preview/single_will")
    assert res.status_code == 200
    unresolved = res.json()["unresolved_placeholders"]
    assert "neverRenderedInFolders" not in unresolved
    assert "[TESTATOR NAME]" in unresolved or "[Client Name]" in unresolved
