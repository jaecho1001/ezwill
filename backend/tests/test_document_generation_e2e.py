"""
E2E test for document generation.
Generates actual DOCX files from mock client data.
Run: .venv/bin/python3 -m pytest tests/test_document_generation_e2e.py -v
"""
import pytest
import os
import sys
import io

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document as DocxDocument
from services.document_generator import DocumentGenerator, resolve_variables

# ── Output directory ─────────────────────────────────────────────────────────

OUTPUT_DIR = "/tmp/ezwill-test-docs"


# ── Mock client data (Kim/Lee dual will precedent) ──────────────────────────

MOCK_VARIABLES = {
    "testatorFullName": "HYUN JUNG KIM",
    "spouseFullName": "MOONYOUNG LEE",
    "city": "City",
    "cityName": "Vaughan",
    "province": "Ontario",
    "primaryExecutorFullName": "MOONYOUNG LEE",
    "backupExecutorFullName": "JOON TAE PARK",
    "childNames": "SUNNY SUN KIM and HELEN SUNWOO KIM",
    "trustDistributionAge": "25",
    "survivalDays": "30",
    "numberOfPages": "12",
    "deponentName": "JAE HYON CHO",
    "deponentCity": "City",
    "deponentCityName": "Toronto",
    "otherWitnessName": "JUNKI LEE",
    "dateOfWill": "November 7, 2023",
    "commissionerName": "",
    "commissionerExpiry": "",
    "willType": "Probate Will",
    "trusteeTitle": "Probate Estate Trustee",
    "otherWillType": "Non-Probate Will",
    "otherTrusteeTitle": "Non-Probate Estate Trustee",
}


# ── Mock clause builders ────────────────────────────────────────────────────

def _build_will_clauses():
    """Build a representative set of will clauses with folders and leaf clauses."""
    return [
        {
            "clause_id": "revocation",
            "isFolder": True,
            "included": True,
            "title": "REVOCATION",
            "sortOrder": 1,
        },
        {
            "clause_id": "revocation_clause",
            "isFolder": False,
            "included": True,
            "title": "Revocation of Prior Wills",
            "templateText": (
                "I, <b>{{testatorFullName}}</b>, of the {{cityName}}, "
                "in the Province of {{province}}, hereby revoke all former Wills "
                "and Codicils made by me."
            ),
            "sortOrder": 2,
        },
        {
            "clause_id": "executors",
            "isFolder": True,
            "included": True,
            "title": "EXECUTORS AND TRUSTEES",
            "sortOrder": 3,
        },
        {
            "clause_id": "primary_executor",
            "isFolder": False,
            "included": True,
            "title": "Appointment of Executor",
            "templateText": (
                "I appoint my spouse, <b>{{primaryExecutorFullName}}</b>, "
                "to be the {{trusteeTitle}} of my estate (my '{{trusteeTitle}}'). "
                "If my spouse is unable or unwilling to act, I appoint "
                "<b>{{backupExecutorFullName}}</b> as alternate {{trusteeTitle}}."
            ),
            "sortOrder": 4,
        },
        {
            "clause_id": "guardianship",
            "isFolder": True,
            "included": True,
            "title": "GUARDIANSHIP",
            "sortOrder": 5,
        },
        {
            "clause_id": "guardian_clause",
            "isFolder": False,
            "included": True,
            "title": "Guardian of Minor Children",
            "templateText": (
                "If my spouse, {{spouseFullName}}, predeceases me or is unable "
                "to act as guardian, I appoint {{backupExecutorFullName}} as "
                "guardian of my minor children, {{childNames}}."
            ),
            "sortOrder": 6,
        },
        {
            "clause_id": "distribution",
            "isFolder": True,
            "included": True,
            "title": "DISTRIBUTION OF ESTATE",
            "sortOrder": 7,
        },
        {
            "clause_id": "spouse_bequest",
            "isFolder": False,
            "included": True,
            "title": "Gift to Spouse",
            "templateText": (
                "I give the entirety of my estate to my spouse, "
                "<b>{{spouseFullName}}</b>, if my spouse survives me by "
                "{{survivalDays}} days."
            ),
            "sortOrder": 8,
        },
        {
            "clause_id": "children_trust",
            "isFolder": False,
            "included": True,
            "title": "Trust for Children",
            "templateText": (
                "If my spouse does not survive me by {{survivalDays}} days, "
                "I direct my {{trusteeTitle}} to divide the residue of my estate "
                "equally among my children, {{childNames}}, to be held in trust "
                "until each child attains the age of {{trustDistributionAge}} years."
            ),
            "sortOrder": 9,
        },
        {
            "clause_id": "survival_clause",
            "isFolder": False,
            "included": True,
            "title": "Survival Requirement",
            "templateText": (
                "A beneficiary must survive me by {{survivalDays}} days to be "
                "entitled to receive any benefit under this Will. If a beneficiary "
                "fails to so survive me, the gift shall be distributed as if that "
                "beneficiary had predeceased me."
            ),
            "sortOrder": 10,
        },
    ]


def _build_short_will_clauses():
    """Build the intentionally small clause set used by the short-form will."""
    return [
        {
            "clause_id": "revocation",
            "isFolder": True,
            "included": True,
            "title": "REVOCATION",
            "sortOrder": 1,
        },
        {
            "clause_id": "revocation_clause",
            "isFolder": False,
            "included": True,
            "title": "Revocation",
            "templateText": (
                "I, {{testatorFullName}}, revoke all former Wills and Codicils "
                "made by me."
            ),
            "sortOrder": 2,
        },
        {
            "clause_id": "executor",
            "isFolder": True,
            "included": True,
            "title": "ESTATE TRUSTEE",
            "sortOrder": 3,
        },
        {
            "clause_id": "executor_clause",
            "isFolder": False,
            "included": True,
            "title": "Appointment",
            "templateText": (
                "I appoint {{primaryExecutorFullName}} to be the Estate Trustee "
                "of this my Will. If {{primaryExecutorFullName}} is unable or "
                "unwilling to act, I appoint {{backupExecutorFullName}} instead."
            ),
            "sortOrder": 4,
        },
        {
            "clause_id": "debts",
            "isFolder": False,
            "included": True,
            "title": "Debts and Taxes",
            "templateText": (
                "I direct my Estate Trustee to pay my just debts, funeral and "
                "testamentary expenses, and taxes payable because of my death."
            ),
            "sortOrder": 5,
        },
        {
            "clause_id": "residue",
            "isFolder": False,
            "included": True,
            "title": "Residue",
            "templateText": (
                "I give the residue of my estate to {{spouseFullName}}, if my "
                "spouse survives me, and otherwise to my children, {{childNames}}, "
                "in equal shares per stirpes."
            ),
            "sortOrder": 6,
        },
        {
            "clause_id": "fla_exclusion",
            "isFolder": False,
            "included": True,
            "title": "Family Law Act Exclusion",
            "templateText": (
                "I declare that property passing under this Will, and income from "
                "it, is excluded from a beneficiary's net family property."
            ),
            "sortOrder": 7,
        },
    ]


def _build_poa_property_clauses():
    """Build POA Property clauses."""
    return [
        {
            "clause_id": "poa_prop_appointment",
            "isFolder": True,
            "included": True,
            "title": "APPOINTMENT",
            "sortOrder": 1,
        },
        {
            "clause_id": "poa_prop_appoint_clause",
            "isFolder": False,
            "included": True,
            "title": "Appointment of Attorney",
            "templateText": (
                "I, <b>{{testatorFullName}}</b>, revoke any previous continuing "
                "power of attorney for property and appoint "
                "<b>{{primaryExecutorFullName}}</b> to be my attorney for property."
            ),
            "sortOrder": 2,
        },
        {
            "clause_id": "poa_prop_authority",
            "isFolder": False,
            "included": True,
            "title": "General Authority",
            "templateText": (
                "I authorize my attorney to do on my behalf anything in respect of "
                "property that I could do if capable of managing property, except "
                "make a will."
            ),
            "sortOrder": 3,
        },
    ]


def _build_poa_personal_care_clauses():
    """Build POA Personal Care clauses."""
    return [
        {
            "clause_id": "poa_pc_appointment",
            "isFolder": True,
            "included": True,
            "title": "APPOINTMENT",
            "sortOrder": 1,
        },
        {
            "clause_id": "poa_pc_appoint_clause",
            "isFolder": False,
            "included": True,
            "title": "Appointment of Attorney for Personal Care",
            "templateText": (
                "I, <b>{{testatorFullName}}</b>, revoke any previous power of "
                "attorney for personal care and appoint "
                "<b>{{primaryExecutorFullName}}</b> to be my attorney for "
                "personal care."
            ),
            "sortOrder": 2,
        },
        {
            "clause_id": "poa_pc_wishes",
            "isFolder": False,
            "included": True,
            "title": "Wishes Regarding Health Care",
            "templateText": (
                "If my attorney is required to make decisions about my health care, "
                "I direct my attorney to consider any wishes I may have expressed."
            ),
            "sortOrder": 3,
        },
    ]


def _build_affidavit_clauses():
    """Build minimal affidavit clauses (the signing page carries the main content)."""
    return [
        {
            "clause_id": "affidavit_intro",
            "isFolder": False,
            "included": True,
            "title": "",
            "templateText": (
                "This Affidavit is made in connection with the "
                "{{willType}} of {{testatorFullName}}, dated {{dateOfWill}}."
            ),
            "sortOrder": 1,
        },
    ]


# ── Fixtures ─────────────────────────────────────────────────────────────────

@pytest.fixture(scope="session", autouse=True)
def setup_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


@pytest.fixture
def generator():
    return DocumentGenerator()


# ── Helper ───────────────────────────────────────────────────────────────────

def _save_and_validate(docx_bytes: bytes, filename: str) -> DocxDocument:
    """Save DOCX bytes to disk and return a parsed Document for assertions."""
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 1024, f"Document too small: {len(docx_bytes)} bytes"

    path = os.path.join(OUTPUT_DIR, filename)
    with open(path, "wb") as f:
        f.write(docx_bytes)

    # Verify parseable
    doc = DocxDocument(io.BytesIO(docx_bytes))
    assert len(doc.paragraphs) > 0, "Document has no paragraphs"
    return doc


def _all_text(doc: DocxDocument) -> str:
    """Extract all text from a DOCX document (paragraphs + tables)."""
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ── Tests ────────────────────────────────────────────────────────────────────

class TestResolveVariables:
    """Test the {{placeholder}} resolution engine."""

    def test_simple_replacement(self):
        result = resolve_variables("Hello {{testatorFullName}}", MOCK_VARIABLES)
        assert result == "Hello HYUN JUNG KIM"

    def test_multiple_replacements(self):
        result = resolve_variables(
            "{{testatorFullName}} of {{cityName}}, {{province}}",
            MOCK_VARIABLES,
        )
        assert result == "HYUN JUNG KIM of Vaughan, Ontario"

    def test_unresolved_becomes_bracket(self):
        result = resolve_variables("Hello {{unknownVar}}", MOCK_VARIABLES)
        assert result == "Hello [unknownVar]"

    def test_empty_input(self):
        assert resolve_variables("", MOCK_VARIABLES) == ""
        assert resolve_variables(None, MOCK_VARIABLES) is None

    def test_no_placeholders(self):
        result = resolve_variables("No placeholders here.", MOCK_VARIABLES)
        assert result == "No placeholders here."


class TestGenerateSingleWill:
    """Generate a single will DOCX (uses probate_will type with 'single_will' alias)."""

    def test_generate_single_will(self, generator):
        clauses = _build_will_clauses()
        docx_bytes = generator.generate_document(
            document_type="probate_will",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = _save_and_validate(docx_bytes, "single_will_kim.docx")
        text = _all_text(doc)
        assert "HYUN JUNG KIM" in text
        assert "MOONYOUNG LEE" in text


class TestGenerateShortFormWill:
    """Generate the compact short-form will without the standard cover page."""

    def test_generate_short_form_will(self, generator):
        clauses = _build_short_will_clauses()
        docx_bytes = generator.generate_document(
            document_type="simple_will_short",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = _save_and_validate(docx_bytes, "simple_will_short_kim.docx")
        text = _all_text(doc)

        assert "SHORT FORM LAST WILL AND TESTAMENT" in text
        assert "HYUN JUNG KIM" in text
        assert "MOONYOUNG LEE" in text
        assert "TESTIMONIUM" in text
        assert "VATURI & CHO LLP" not in text
        assert "SCHEDULE" not in text


class TestGenerateProbateWill:
    """Generate a full probate will with all Tier 2 clauses."""

    def test_generate_probate_will(self, generator):
        clauses = _build_will_clauses()
        docx_bytes = generator.generate_document(
            document_type="probate_will",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = _save_and_validate(docx_bytes, "probate_will_kim.docx")
        text = _all_text(doc)

        # Cover page content
        assert "VATURI & CHO LLP" in text
        assert "LAST WILL AND TESTAMENT" in text

        # Clause headings
        assert "REVOCATION" in text
        assert "EXECUTORS AND TRUSTEES" in text
        assert "GUARDIANSHIP" in text
        assert "DISTRIBUTION OF ESTATE" in text

        # Variable resolution in clause body
        assert "HYUN JUNG KIM" in text
        assert "MOONYOUNG LEE" in text
        assert "JOON TAE PARK" in text
        assert "SUNNY SUN KIM and HELEN SUNWOO KIM" in text

        # Signing page content (table-based)
        assert "TESTIMONIUM" in text

        # Schedule A for probate will
        assert "SCHEDULE" in text
        assert "EXCLUDED PROPERTY" in text


class TestGenerateNonProbateWill:
    """Generate a non-probate will DOCX."""

    def test_generate_non_probate_will(self, generator):
        np_variables = {**MOCK_VARIABLES, "trusteeTitle": "Non-Probate Estate Trustee"}
        clauses = _build_will_clauses()
        docx_bytes = generator.generate_document(
            document_type="non_probate_will",
            clauses=clauses,
            variables=np_variables,
        )
        doc = _save_and_validate(docx_bytes, "non_probate_will_kim.docx")
        text = _all_text(doc)
        assert "NON-PROBATE" in text
        assert "HYUN JUNG KIM" in text
        assert "Non-Probate Estate Trustee" in text


class TestGenerateAffidavit:
    """Generate an affidavit of execution DOCX."""

    def test_generate_affidavit(self, generator):
        clauses = _build_affidavit_clauses()
        docx_bytes = generator.generate_document(
            document_type="affidavit_execution",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = _save_and_validate(docx_bytes, "affidavit_execution_kim.docx")
        text = _all_text(doc)
        assert "AFFIDAVIT OF EXECUTION" in text
        assert "JAE HYON CHO" in text
        assert "HYUN JUNG KIM" in text
        # Exhibit A
        assert 'EXHIBIT "A"' in text or "EXHIBIT" in text


class TestGeneratePOAProperty:
    """Generate POA Property DOCX."""

    def test_generate_poa_property(self, generator):
        clauses = _build_poa_property_clauses()
        docx_bytes = generator.generate_document(
            document_type="poa_property",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = _save_and_validate(docx_bytes, "poa_property_kim.docx")
        text = _all_text(doc)
        assert "CONTINUING POWER OF ATTORNEY FOR PROPERTY" in text
        assert "HYUN JUNG KIM" in text
        assert "MOONYOUNG LEE" in text
        assert "EXECUTION" in text


class TestGeneratePOAPersonalCare:
    """Generate POA Personal Care DOCX."""

    def test_generate_poa_personal_care(self, generator):
        clauses = _build_poa_personal_care_clauses()
        docx_bytes = generator.generate_document(
            document_type="poa_personal_care",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = _save_and_validate(docx_bytes, "poa_personal_care_kim.docx")
        text = _all_text(doc)
        assert "POWER OF ATTORNEY FOR PERSONAL CARE" in text
        assert "HYUN JUNG KIM" in text


class TestSigningPageHasTable:
    """Verify the signing page uses Word tables (not just plain text)."""

    def test_signing_page_has_table(self, generator):
        clauses = _build_will_clauses()
        docx_bytes = generator.generate_document(
            document_type="probate_will",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = DocxDocument(io.BytesIO(docx_bytes))
        # The document should have tables (attestation + witness blocks)
        assert len(doc.tables) >= 2, (
            f"Expected at least 2 tables for signing page, found {len(doc.tables)}"
        )


class TestCoverPage:
    """Verify the cover page has firm name and document title."""

    def test_cover_page_firm_name(self, generator):
        clauses = _build_will_clauses()
        docx_bytes = generator.generate_document(
            document_type="probate_will",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = DocxDocument(io.BytesIO(docx_bytes))
        # First paragraphs should be firm name and title
        first_texts = [p.text for p in doc.paragraphs[:10]]
        combined = " ".join(first_texts)
        assert "VATURI & CHO LLP" in combined
        assert "LAST WILL AND TESTAMENT" in combined

    def test_cover_page_client_name(self, generator):
        clauses = _build_will_clauses()
        docx_bytes = generator.generate_document(
            document_type="probate_will",
            clauses=clauses,
            variables=MOCK_VARIABLES,
        )
        doc = DocxDocument(io.BytesIO(docx_bytes))
        first_texts = [p.text for p in doc.paragraphs[:10]]
        combined = " ".join(first_texts)
        assert "HYUN JUNG KIM" in combined


class TestGenerateAllDocuments:
    """Generate all documents for a dual will client."""

    def test_generate_all_documents(self, generator):
        clause_selections = {
            "probate_will": _build_will_clauses(),
            "non_probate_will": _build_will_clauses(),
            "poa_property": _build_poa_property_clauses(),
            "poa_personal_care": _build_poa_personal_care_clauses(),
            "affidavit_execution": _build_affidavit_clauses(),
            "affidavit_execution_np": _build_affidavit_clauses(),
            "affidavit_execution_poa_prop": _build_affidavit_clauses(),
            "affidavit_execution_poa_pc": _build_affidavit_clauses(),
        }

        results = generator.generate_all_documents(
            draft_data={"id": "test-draft-001", "client_name": "HYUN JUNG KIM"},
            clause_selections=clause_selections,
            variables=MOCK_VARIABLES,
        )

        # Verify we get all 8 document types
        assert len(results) == 8, (
            f"Expected 8 documents, got {len(results)}: {list(results.keys())}"
        )

        expected_types = {
            "probate_will",
            "non_probate_will",
            "poa_property",
            "poa_personal_care",
            "affidavit_execution",
            "affidavit_execution_np",
            "affidavit_execution_poa_prop",
            "affidavit_execution_poa_pc",
        }
        assert set(results.keys()) == expected_types

        # Save all and verify each
        for doc_type, docx_bytes in results.items():
            filename = f"all_{doc_type}_kim.docx"
            doc = _save_and_validate(docx_bytes, filename)
            text = _all_text(doc)
            assert "VATURI & CHO LLP" in text, f"{doc_type} missing firm name"
            assert "HYUN JUNG KIM" in text, f"{doc_type} missing testator name"
