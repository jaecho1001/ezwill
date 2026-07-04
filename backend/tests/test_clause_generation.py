"""Regression tests for the P0 document-generation fix.

Root cause: ew_clause_selections stored only custom_text, so clauses the lawyer
never hand-edited reached the generator with an empty body. The generator reads
template_text (snake_case, as stored in the DB row), so once we persist
template_text a complete will renders. These tests pin that behaviour using the
exact snake_case shape a DB row now carries.
"""

from __future__ import annotations

import io
import os
import sys

from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from models import ClauseSelection
from services.document_generator import (
    DocumentGenerator,
    map_people_to_variables,
    vault_to_variables,
)


VARIABLES = {
    "testatorFullName": "HYUN JUNG KIM",
    "city": "City",
    "cityName": "Vaughan",
    "province": "Ontario",
    "primaryExecutorFullName": "MOONYOUNG LEE",
    "documentDate": "April 8, 2026",
}


def _docx_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def _db_row_clauses() -> list[dict]:
    """Clauses shaped exactly as get_clause_selections returns them after the
    migration — snake_case template_text, no custom_text (lawyer never edited)."""
    return [
        {
            "clause_id": "rev",
            "title": "Revocation",
            "template_text": "",
            "custom_text": None,
            "is_folder": True,
            "included": True,
            "sort_order": 1,
        },
        {
            "clause_id": "rev-single",
            "title": "Revocation of Prior Wills",
            "template_text": (
                "I, {{testatorFullName}}, of the {{city}} of {{cityName}}, in the "
                "Province of {{province}}, revoke all former Wills and declare this "
                "to be my Last Will and Testament."
            ),
            "custom_text": None,
            "is_folder": False,
            "included": True,
            "sort_order": 2,
        },
        {
            "clause_id": "appt-primary",
            "title": "Appointment of Estate Trustee",
            "template_text": "I appoint {{primaryExecutorFullName}} as the Estate Trustee of this my Will.",
            "custom_text": None,
            "is_folder": False,
            "included": True,
            "sort_order": 3,
        },
    ]


def test_generator_renders_template_text_from_db_row_shape():
    """The exact regression: a clause with only template_text (no custom_text)
    must render its resolved body, not an empty paragraph."""
    docx_bytes = DocumentGenerator().generate_document(
        document_type="single_will",
        clauses=_db_row_clauses(),
        variables=VARIABLES,
    )
    text = _docx_text(docx_bytes)

    # Body text is present and variables are resolved.
    assert "revoke all former Wills" in text
    assert "HYUN JUNG KIM" in text
    assert "MOONYOUNG LEE" in text
    assert "Vaughan" in text
    # No unresolved placeholders leaked into the document.
    assert "{{" not in text and "}}" not in text


def test_custom_text_overrides_template_text():
    clauses = _db_row_clauses()
    clauses[1]["custom_text"] = "I revoke everything — custom lawyer wording for {{testatorFullName}}."
    text = _docx_text(
        DocumentGenerator().generate_document(
            document_type="single_will", clauses=clauses, variables=VARIABLES
        )
    )
    assert "custom lawyer wording for HYUN JUNG KIM" in text
    assert "revoke all former Wills" not in text  # template body was overridden


def test_empty_template_text_is_the_bug_we_fixed():
    """Sanity check that proves the test is meaningful: with NO text at all
    (the pre-fix DB shape), the body clause contributes no resolved sentence."""
    clauses = _db_row_clauses()
    for c in clauses:
        c["template_text"] = ""
        c["custom_text"] = None
    text = _docx_text(
        DocumentGenerator().generate_document(
            document_type="single_will", clauses=clauses, variables=VARIABLES
        )
    )
    assert "revoke all former Wills" not in text  # empty body — the P0 symptom


def test_people_roles_map_to_variables_using_stored_vocabulary():
    """Regression: the frontend stores role 'executor'/'attorney_property'/
    'attorney_care', but the generator used to look for 'primary_executor'/
    'poa_*', so executor & attorney names never rendered."""
    people = [
        {"role": "executor", "first_name": "Moonyoung", "last_name": "Lee"},
        {"role": "backup_executor", "first_name": "Joon", "last_name": "Park"},
        {"role": "attorney_property", "first_name": "Ava", "last_name": "Kim"},
        {"role": "attorney_care", "first_name": "Ben", "last_name": "Cho"},
        {"role": "guardian", "first_name": "Sung", "last_name": "Hee"},
        {"role": "beneficiary", "first_name": "Ignore", "last_name": "Me"},
    ]
    v = map_people_to_variables(people)
    assert v["primaryExecutorFullName"] == "Moonyoung Lee"
    assert v["backupExecutorFullName"] == "Joon Park"
    assert v["poaPropertyAttorneyFullName"] == "Ava Kim"
    assert v["poaPersonalCareAttorneyFullName"] == "Ben Cho"
    assert v["guardianFullName"] == "Sung Hee"
    # beneficiaries aren't name-variable roles
    assert "beneficiary" not in str(v)


def test_people_roles_accept_legacy_vocabulary_too():
    v = map_people_to_variables([
        {"role": "primary_executor", "first_name": "Old", "last_name": "Name"},
        {"role": "poa_property_attorney", "first_name": "Legacy", "last_name": "Attorney"},
    ])
    assert v["primaryExecutorFullName"] == "Old Name"
    assert v["poaPropertyAttorneyFullName"] == "Legacy Attorney"


def test_build_variables_reads_questionnaire_sections():
    """P0 #2: section answers (about_you/your_family/your_estate) must flow into
    the generator variables. Before migration 29 these columns didn't exist, so
    they never reached _build_variables."""
    from routes.documents import _build_variables

    draft = {
        "client_first_name": "Jane",
        "client_last_name": "Doe",
        "province": "ON",
        "about_you": {"city": "Vaughan", "maritalStatus": "married"},
        "your_family": {
            "spouseFullName": "Moonyoung Lee",
            "children": [
                {"firstName": "Sunny", "lastName": "Kim"},
                {"firstName": "Helen", "lastName": "Kim"},
            ],
        },
        "your_estate": {"survivalDays": 30, "trustDistributionAge": 25},
        "people": [{"role": "executor", "first_name": "Moonyoung", "last_name": "Lee"}],
    }
    v = _build_variables(draft)
    assert v["testatorFullName"] == "Jane Doe"
    assert v["cityName"] == "Vaughan"
    assert v["spouseFullName"] == "Moonyoung Lee"
    assert "Sunny Kim" in v["childNames"] and "Helen Kim" in v["childNames"]
    assert v["trustDistributionAge"] == "25"
    assert v["primaryExecutorFullName"] == "Moonyoung Lee"


def test_vault_to_variables_projects_intake_data():
    """P0 #2 (vault): conversational-intake facts must project into the same
    {{placeholder}} variables the generator resolves."""
    vault = {
        "testator": {"fullName": "Jane Doe", "address": "10 King St, Vaughan, ON"},
        "spouse": {"included": True, "fullName": "Moonyoung Lee"},
        "children": [{"id": "1", "fullName": "Sunny Kim"}, {"id": "2", "fullName": "Helen Kim"}],
        "executors": [
            {"id": "e1", "fullName": "Moonyoung Lee", "isBackup": False},
            {"id": "e2", "fullName": "Joon Park", "isBackup": True},
        ],
        "guardians": [{"id": "g1", "fullName": "Sung Hee", "isBackup": False}],
    }
    v = vault_to_variables(vault)
    assert v["testatorFullName"] == "Jane Doe"
    assert v["cityName"] == "Vaughan"
    assert v["province"] == "Ontario"
    assert v["spouseFullName"] == "Moonyoung Lee"
    assert v["childNames"] == "Sunny Kim and Helen Kim"
    assert v["primaryExecutorFullName"] == "Moonyoung Lee"
    assert v["backupExecutorFullName"] == "Joon Park"
    assert v["guardianFullName"] == "Sung Hee"


def test_vault_to_variables_empty_is_noop():
    assert vault_to_variables({}) == {}
    assert vault_to_variables(None) == {}
    # An uninvolved spouse is not projected.
    assert "spouseFullName" not in vault_to_variables(
        {"testator": {}, "spouse": {"included": False, "fullName": "X"},
         "children": [], "executors": [], "guardians": []}
    )


def test_build_variables_overlays_vault_when_present():
    """A chat-only client (no questionnaire sections) still gets their data in
    the document via the vault overlay."""
    from routes.documents import _build_variables

    draft = {
        "client_first_name": "",
        "client_last_name": "",
        "province": "ON",
        "vault": {
            "testator": {"fullName": "Chat Only", "address": "5 Yonge St, Toronto, ON"},
            "children": [],
            "executors": [{"id": "e", "fullName": "Trusted Friend", "isBackup": False}],
            "guardians": [],
        },
    }
    v = _build_variables(draft)
    assert v["testatorFullName"] == "Chat Only"
    assert v["cityName"] == "Toronto"
    assert v["primaryExecutorFullName"] == "Trusted Friend"


def test_clause_body_preserves_sub_item_structure_from_raw_text():
    """P1 fidelity: (a)/(i) hanging-indent structure must survive to Word instead
    of collapsing into one run-on paragraph."""
    from services.document_generator import clause_body_to_blocks

    raw = (
        "I hold the residue on the following terms:\n\n"
        "(a) to pay income to my spouse;\n\n"
        "(b) then to divide the capital as follows:\n\n"
        "(i) one half to my children;\n\n"
        "(ii) one half to charity."
    )
    blocks = clause_body_to_blocks(raw)
    assert [b["indent"] for b in blocks] == [1, 2, 2, 3, 3]
    assert [b["marker"] for b in blocks] == [None, "(a)", "(b)", "(i)", "(ii)"]

    docx_bytes = DocumentGenerator().generate_document(
        "single_will",
        [{"clause_id": "x", "title": "Residue", "template_text": raw,
          "is_folder": False, "included": True, "sort_order": 1}],
        {},
    )
    d = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "(a)\t" in text and "(i)\t" in text  # markers rendered as literal text
    a_para = next(p for p in d.paragraphs if p.text.startswith("(a)"))
    i_para = next(p for p in d.paragraphs if p.text.startswith("(i)"))
    assert round(a_para.paragraph_format.left_indent.inches, 2) == 1.0
    assert round(a_para.paragraph_format.first_line_indent.inches, 2) == -0.5
    assert round(i_para.paragraph_format.left_indent.inches, 2) == 1.5


def test_clause_body_html_preserves_blocks_and_inline_italic():
    from services.document_generator import clause_body_to_blocks

    html = (
        '<p data-indent="1">The <i>Family Law Act</i> applies:</p>'
        '<p data-indent="2" data-marker="(a)">the lettered item.</p>'
    )
    blocks = clause_body_to_blocks(html)
    assert [b["indent"] for b in blocks] == [1, 2]
    assert blocks[1]["marker"] == "(a)"
    assert any(r[2] for r in blocks[0]["runs"])  # an italic run survived


def test_clause_selection_model_accepts_template_fields():
    sel = ClauseSelection(
        clause_id="rev-single",
        template_text="I, {{testatorFullName}}, revoke all former Wills.",
        title="Revocation of Prior Wills",
        is_folder=False,
        sort_order=2,
    )
    assert sel.template_text.startswith("I, {{testatorFullName}}")
    assert sel.title == "Revocation of Prior Wills"
    assert sel.is_folder is False
    # Backwards compatible: the old minimal shape still validates.
    legacy = ClauseSelection(clause_id="x")
    assert legacy.template_text is None and legacy.is_folder is False
