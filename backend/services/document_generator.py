"""
EZWill Document Generation Service
Generates professional DOCX documents from clause selections using python-docx.
Supports all built-in document types with table-based signing pages.
"""

from __future__ import annotations

import re
import io
import logging
from html.parser import HTMLParser
from datetime import date
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# ── Constants ───────────────────────────────────────────────────────────────

FIRM_NAME = "VATURI & CHO LLP"
FIRM_ADDRESS_LINE1 = "1110 Finch Avenue West, Suite 310"
FIRM_ADDRESS_LINE2 = "Toronto, ON  M3J 2T2"

DOCUMENT_TITLES: dict[str, str] = {
    "simple_will_short": "SHORT FORM LAST WILL AND TESTAMENT",
    "single_will": "LAST WILL AND TESTAMENT",
    "probate_will": "LAST WILL AND TESTAMENT (PROBATE WILL)",
    "non_probate_will": "LAST WILL AND TESTAMENT (NON-PROBATE WILL)",
    "poa_property": "CONTINUING POWER OF ATTORNEY FOR PROPERTY",
    "poa_personal_care": "POWER OF ATTORNEY FOR PERSONAL CARE",
    "affidavit_execution": "AFFIDAVIT OF EXECUTION",
    "affidavit_execution_probate": "AFFIDAVIT OF EXECUTION (PROBATE WILL)",
    "affidavit_execution_non_probate": "AFFIDAVIT OF EXECUTION (NON-PROBATE WILL)",
}

WILL_TYPES = {"simple_will_short", "single_will", "probate_will", "non_probate_will"}
COMPACT_WILL_TYPES = {"simple_will_short"}
POA_TYPES = {"poa_property", "poa_personal_care"}
AFFIDAVIT_TYPES = {
    "affidavit_execution", "affidavit_execution_probate",
    "affidavit_execution_non_probate",
}

VARIABLE_PATTERN = re.compile(r"\{\{(\w+)\}\}")


# ── HTML Parser ─────────────────────────────────────────────────────────────

class _SimpleHTMLParser(HTMLParser):
    """Parses simple HTML and produces a list of (text, bold, italic, underline) runs."""

    def __init__(self):
        super().__init__()
        self.runs: list[tuple[str, bool, bool, bool]] = []
        self._bold = False
        self._italic = False
        self._underline = False

    def handle_starttag(self, tag: str, attrs):
        tag = tag.lower()
        if tag in ("b", "strong"):
            self._bold = True
        elif tag in ("i", "em"):
            self._italic = True
        elif tag == "u":
            self._underline = True

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if tag in ("b", "strong"):
            self._bold = False
        elif tag in ("i", "em"):
            self._italic = False
        elif tag == "u":
            self._underline = False

    def handle_data(self, data: str):
        if data:
            self.runs.append((data, self._bold, self._italic, self._underline))


# ── Helpers ─────────────────────────────────────────────────────────────────

# Maps a stored ew_people.role to the {{placeholder}} it fills. The keys are the
# roles the frontend actually stores (see extractPeople + the ew_people_role_check
# constraint: executor / attorney_property / attorney_care / …). The older
# generator vocabulary (primary_executor / poa_*_attorney) is accepted too so a
# mix of legacy rows still resolves. This is the single source of truth for both
# the DOCX generator and the client review portal.
_ROLE_TO_VARIABLE = {
    "executor": "primaryExecutorFullName",
    "primary_executor": "primaryExecutorFullName",
    "backup_executor": "backupExecutorFullName",
    "attorney_property": "poaPropertyAttorneyFullName",
    "poa_property_attorney": "poaPropertyAttorneyFullName",
    "attorney_care": "poaPersonalCareAttorneyFullName",
    "poa_personal_care_attorney": "poaPersonalCareAttorneyFullName",
    "backup_attorney": "backupAttorneyFullName",
    "guardian": "guardianFullName",
    "backup_guardian": "backupGuardianFullName",
}


def map_people_to_variables(people: list) -> dict:
    """Resolve a draft's people list into executor/attorney/guardian name
    variables. Person name fields are the snake_case columns stored by
    upsert_people (first_name/last_name). First non-empty name per role wins."""
    out: dict = {}
    for person in people or []:
        var = _ROLE_TO_VARIABLE.get(person.get("role", ""))
        if not var or out.get(var):
            continue
        full_name = f"{person.get('first_name', '')} {person.get('last_name', '')}".strip()
        if full_name:
            out[var] = full_name
    return out


_PROVINCE_RE = re.compile(
    r"\b(Ontario|ON|Quebec|QC|British Columbia|BC|Alberta|AB|Manitoba|MB|"
    r"Saskatchewan|SK|Nova Scotia|NS|New Brunswick|NB|Newfoundland|NL|"
    r"Prince Edward Island|PE|PEI)\b",
    re.IGNORECASE,
)


def _extract_city(address: str) -> str:
    parts = [p.strip() for p in address.split(",") if p.strip()]
    if len(parts) >= 2:
        return parts[-2]
    return parts[0] if parts else ""


def _extract_province(address: str) -> str:
    m = _PROVINCE_RE.search(address)
    if not m:
        return ""
    raw = m.group(1)
    if raw.upper() == "ON" or raw.lower() == "ontario":
        return "Ontario"
    return raw


def _list_names(names: list) -> str:
    clean = [n for n in names if n]
    if not clean:
        return ""
    if len(clean) == 1:
        return clean[0]
    if len(clean) == 2:
        return f"{clean[0]} and {clean[1]}"
    return f"{', '.join(clean[:-1])}, and {clean[-1]}"


def vault_to_variables(vault: dict) -> dict:
    """Project a conversational-intake WillVault into {{placeholder}} variables.

    Python mirror of the frontend vaultToVariables (src/lib/will-documents/
    vault-to-variables.ts). Only emits variables backed by real vault data, so
    merging it over questionnaire-derived variables never clobbers a filled
    field with a blank. Keep the two in sync when the vault schema changes.
    """
    if not vault:
        return {}
    v: dict = {}

    testator = vault.get("testator") or {}
    if testator.get("fullName"):
        v["testatorFullName"] = testator["fullName"]
    if testator.get("address"):
        city = _extract_city(testator["address"])
        if city:
            v["city"] = city
            v["cityName"] = city
        province = _extract_province(testator["address"])
        if province:
            v["province"] = province

    spouse = vault.get("spouse") or {}
    if spouse.get("included") and spouse.get("fullName"):
        v["spouseFullName"] = spouse["fullName"]

    child_names = [c.get("fullName") for c in (vault.get("children") or [])]
    names = _list_names(child_names)
    if names:
        v["childNames"] = names

    executors = vault.get("executors") or []
    primary = next((e for e in executors if not e.get("isBackup")), None)
    backup = next((e for e in executors if e.get("isBackup")), None)
    if primary and primary.get("fullName"):
        v["primaryExecutorFullName"] = primary["fullName"]
    if backup and backup.get("fullName"):
        v["backupExecutorFullName"] = backup["fullName"]

    primary_guardian = next(
        (g for g in (vault.get("guardians") or []) if not g.get("isBackup")), None
    )
    if primary_guardian and primary_guardian.get("fullName"):
        v["guardianFullName"] = primary_guardian["fullName"]
        v["primaryGuardianFullName"] = primary_guardian["fullName"]

    return v


def firm_variables(settings: dict) -> dict:
    """Project saved firm settings into document variables (firm identity that
    the cover page + review portal render). Keys the frontend settings page
    wrote in camelCase (firmName, address1, lsoNumber, ...)."""
    if not settings:
        return {}
    firm = settings.get("firm") or {}
    v: dict = {}
    if firm.get("firmName"):
        v["firmName"] = firm["firmName"]
    line1 = ", ".join(p for p in (firm.get("address1"), firm.get("address2")) if p)
    if line1:
        v["firmAddressLine1"] = line1
    line2 = ", ".join(p for p in (firm.get("city"), firm.get("province")) if p)
    if firm.get("postalCode"):
        line2 = (line2 + "  " + firm["postalCode"]).strip()
    if line2:
        v["firmAddressLine2"] = line2
    if firm.get("lsoNumber"):
        v["lsoNumber"] = firm["lsoNumber"]
    return v


def resolve_variables(text: str, variables: dict) -> str:
    """
    Replace {{variableName}} placeholders with values from variables dict.
    Unresolved placeholders become [variableName].
    """
    if not text:
        return text

    def _replace(match: re.Match) -> str:
        key = match.group(1)
        value = variables.get(key)
        if value is not None:
            return str(value)
        # Try camelCase -> snake_case lookup
        snake_key = re.sub(r"([A-Z])", r"_\1", key).lower().lstrip("_")
        value = variables.get(snake_key)
        if value is not None:
            return str(value)
        return f"[{key}]"

    return VARIABLE_PATTERN.sub(_replace, text)


def html_to_docx_runs(paragraph, html_text: str) -> None:
    """
    Parse simple HTML (<b>, <i>, <u>, <strong>, <em>) and add formatted runs
    to a python-docx paragraph. Falls back to plain text if no HTML tags found.
    """
    if not html_text:
        return

    if "<" not in html_text:
        paragraph.add_run(html_text)
        return

    parser = _SimpleHTMLParser()
    try:
        parser.feed(html_text)
    except Exception:
        paragraph.add_run(html_text)
        return

    if not parser.runs:
        paragraph.add_run(html_text)
        return

    for text, bold, italic, underline in parser.runs:
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if underline:
            run.underline = True


def _strip_html(text: str) -> str:
    """Remove HTML tags, returning plain text."""
    if not text or "<" not in text:
        return text or ""
    return re.sub(r"<[^>]+>", "", text)


# ── Structured clause rendering ─────────────────────────────────────────────
# A clause body carries hierarchy the flat run-based renderer lost: an intro
# paragraph (level 1), lettered sub-items (a)/(b) (level 2), and roman
# sub-sub-items (i)/(ii) (level 3), each hanging-indented under its marker.
# It arrives in two shapes and both must render identically to the editor:
#   - lawyer-edited  -> HTML: <p data-indent="2" data-marker="(a)">text</p>
#   - untouched      -> raw template_text with literal "(a)"/"(i)" + blank lines
# We normalize both into "blocks" ({indent, marker, align, runs, heading}) and
# lay each out as its own docx paragraph. Indent levels mirror globals.css:
# 3.5em/7em/10.5em tab stops == 0.5in per level, hanging indent 0.5in.

_ROMAN_SEQUENCE = [
    "i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x",
    "xi", "xii", "xiii", "xiv", "xv", "xvi", "xvii", "xviii", "xix", "xx",
]


def _ends_with_colon(s: str) -> bool:
    return bool(re.search(r":\s*$", s))


def _inline_runs(text: str) -> list:
    """Split a paragraph's (possibly inline-HTML) text into (text, b, i, u) runs."""
    if not text:
        return []
    if "<" not in text:
        return [(text, False, False, False)]
    parser = _SimpleHTMLParser()
    try:
        parser.feed(text)
    except Exception:
        return [(_strip_html(text), False, False, False)]
    return parser.runs or [(_strip_html(text), False, False, False)]


def _format_clause_text(raw: str) -> list:
    """Python port of the frontend formatClauseText (format-clause.ts): split
    raw clause text into paragraphs with indent level + auto/explicit markers.
    Kept in sync with the editor so the .docx matches the on-screen preview."""
    if not raw:
        return []
    use_double = "\n\n" in raw
    parts = re.split(r"\n\n+" if use_double else r"\n+", raw)
    parts = [p.strip() for p in parts if p.strip()]
    if not parts:
        return []

    result: list = []
    mode = "body"
    letter_counter = 0
    roman_counter = 0

    for i, text in enumerate(parts):
        roman_match = re.match(r"^\(([ivx]+)\)\s+([\s\S]*)", text)
        letter_match = re.match(r"^\(([a-z])\)\s+([\s\S]*)", text)

        if roman_match:
            body = roman_match.group(2).strip()
            result.append({"text": body, "indent": 3, "marker": f"({roman_match.group(1)})"})
            mode = "romans"
            idx = _ROMAN_SEQUENCE.index(roman_match.group(1)) if roman_match.group(1) in _ROMAN_SEQUENCE else -1
            roman_counter = idx + 1 if idx >= 0 else roman_counter + 1
            continue

        if letter_match:
            body = letter_match.group(2).strip()
            result.append({"text": body, "indent": 2, "marker": f"({letter_match.group(1)})"})
            letter_counter = ord(letter_match.group(1)) - 0x61 + 1
            mode = "romans" if _ends_with_colon(body) else "letters"
            if mode == "romans":
                roman_counter = 0
            continue

        if i == 0:
            result.append({"text": text, "indent": 1, "marker": None})
            if _ends_with_colon(text):
                mode = "letters"
                letter_counter = 0
                roman_counter = 0
            continue

        if mode == "romans":
            marker = f"({_ROMAN_SEQUENCE[roman_counter]})" if roman_counter < len(_ROMAN_SEQUENCE) else None
            if marker:
                roman_counter += 1
            result.append({"text": text, "indent": 3, "marker": marker})
        elif mode == "letters":
            marker = f"({chr(0x61 + letter_counter)})"
            letter_counter += 1
            result.append({"text": text, "indent": 2, "marker": marker})
            if _ends_with_colon(text):
                mode = "romans"
                roman_counter = 0
        else:
            result.append({"text": text, "indent": 1, "marker": None})

    return result


class _BlockHTMLParser(HTMLParser):
    """Parse editor HTML into block paragraphs, preserving data-indent /
    data-marker / data-num and inline b/i/u runs."""

    _BLOCK_TAGS = ("p", "h3", "h4", "li")

    def __init__(self):
        super().__init__()
        self.blocks: list = []
        self._cur = None
        self._bold = False
        self._italic = False
        self._underline = False

    def _open(self, indent, marker, align, heading):
        self._cur = {
            "indent": indent, "marker": marker, "align": align,
            "heading": heading, "runs": [],
        }

    def _close(self):
        if self._cur is not None:
            if self._cur["runs"] or self._cur["marker"]:
                self.blocks.append(self._cur)
            self._cur = None

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in self._BLOCK_TAGS:
            self._close()
            d = dict(attrs)
            try:
                indent = int(d.get("data-indent") or 1)
            except (TypeError, ValueError):
                indent = 1
            indent = min(max(indent, 1), 3)
            marker = d.get("data-marker") or d.get("data-num") or None
            style = (d.get("style") or "").lower()
            align = "center" if "center" in style else ("right" if "right" in style else None)
            self._open(indent, marker, align, heading=tag in ("h3", "h4"))
        elif tag in ("b", "strong"):
            self._bold = True
        elif tag in ("i", "em"):
            self._italic = True
        elif tag == "u":
            self._underline = True
        elif tag == "br" and self._cur is not None:
            self._cur["runs"].append(("\n", self._bold, self._italic, self._underline))

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in self._BLOCK_TAGS:
            self._close()
        elif tag in ("b", "strong"):
            self._bold = False
        elif tag in ("i", "em"):
            self._italic = False
        elif tag == "u":
            self._underline = False

    def handle_data(self, data):
        if not data:
            return
        if self._cur is None:
            self._open(1, None, None, heading=False)
        self._cur["runs"].append((data, self._bold, self._italic, self._underline))


def clause_body_to_blocks(text: str):
    """Normalize a clause body (edited HTML or raw template text) into layout
    blocks. Returns None on parse failure so the caller can fall back to plain."""
    if not text:
        return []
    if re.search(r"<\s*(p|h3|h4|li|ul|ol|div)\b", text, re.IGNORECASE):
        parser = _BlockHTMLParser()
        try:
            parser.feed(text)
        except Exception:
            return None
        parser._close()
        return parser.blocks
    # Raw text: structure via the ported formatter, then inline runs per line.
    blocks = []
    for para in _format_clause_text(text):
        blocks.append({
            "indent": para["indent"], "marker": para.get("marker"),
            "align": None, "heading": False, "runs": _inline_runs(para["text"]),
        })
    return blocks


# ── Document Generator ──────────────────────────────────────────────────────

class DocumentGenerator:
    """
    Generates professional DOCX documents from clause selections.
    Supports all built-in Ontario estate-planning document types with
    table-based signing pages.
    """

    def generate_document(
        self,
        document_type: str,
        clauses: list[dict],
        variables: dict,
        signing_data: Optional[dict] = None,
    ) -> bytes:
        """
        Generate a DOCX file from clause selections and return it as bytes.

        Args:
            document_type: One of the supported document types.
            clauses: List of clause dicts with keys: clause_id, included,
                     templateText, customText, sortOrder, isFolder, title.
            variables: Dict of placeholder variables to resolve.
            signing_data: Optional dict with signing-page specifics
                          (witness names, commissioner, etc.).

        Returns:
            The generated DOCX file content as bytes.
        """
        doc = Document()
        compact = document_type in COMPACT_WILL_TYPES
        self._set_document_styles(doc, compact=compact)
        if compact:
            self._create_compact_will_heading(doc, document_type, variables)
        else:
            self._create_cover_page(doc, document_type, variables)

        # Filter and sort clauses
        included = [c for c in clauses if c.get("included", True)]
        included.sort(key=lambda c: c.get("sortOrder", c.get("sort_order", 0)))

        clause_number = 0
        for clause in included:
            is_folder = clause.get("isFolder", clause.get("is_folder", False))
            if is_folder:
                self._add_folder_heading(doc, clause, variables)
            else:
                clause_number += 1
                self._add_clause(doc, clause, variables, clause_number)

        # Signing pages
        if signing_data or document_type in WILL_TYPES | POA_TYPES | AFFIDAVIT_TYPES:
            sd = signing_data or {}
            self._create_signing_page(
                doc, document_type, sd, variables, page_break=not compact
            )

        # Schedule A for probate will
        if document_type == "probate_will":
            self._create_schedule_a(doc, variables)

        self._add_page_numbers(doc)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Cover Page ──────────────────────────────────────────────────────────

    def _create_cover_page(
        self, doc: Document, document_type: str, variables: dict
    ) -> None:
        """Add a firm-branded cover page."""
        # Firm identity from saved settings, falling back to the built-in default.
        firm_name = variables.get("firmName") or FIRM_NAME
        addr1 = variables.get("firmAddressLine1") or FIRM_ADDRESS_LINE1
        addr2 = variables.get("firmAddressLine2") or FIRM_ADDRESS_LINE2

        # Firm name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(firm_name)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

        # Address
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(addr1)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(addr2)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

        doc.add_paragraph()  # spacer

        # Document title
        title = DOCUMENT_TITLES.get(document_type, document_type.upper().replace("_", " "))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

        doc.add_paragraph()  # spacer

        # Client name
        client_name = (
            variables.get("testatorFullName")
            or variables.get("testator_full_name")
            or variables.get("clientFullName")
            or variables.get("client_full_name")
            or "[Client Name]"
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"of  {client_name}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        doc.add_paragraph()  # spacer

        # Date
        doc_date = variables.get("documentDate") or date.today().strftime("%B %d, %Y")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(doc_date)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        # Page break after cover
        doc.add_page_break()

    def _create_compact_will_heading(
        self, doc: Document, document_type: str, variables: dict
    ) -> None:
        """Add an inline heading for short-form wills instead of a cover page."""
        title = DOCUMENT_TITLES.get(document_type, document_type.upper().replace("_", " "))
        client_name = (
            variables.get("testatorFullName")
            or variables.get("testator_full_name")
            or variables.get("clientFullName")
            or variables.get("client_full_name")
            or "[Client Name]"
        )
        doc_date = variables.get("documentDate") or date.today().strftime("%B %d, %Y")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(client_name).upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(doc_date))
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    # ── Folder Heading ──────────────────────────────────────────────────────

    def _add_folder_heading(
        self, doc: Document, clause: dict, variables: dict
    ) -> None:
        """Add a bold, caps folder heading (e.g. 'EXECUTORS AND TRUSTEES')."""
        title = clause.get("title", clause.get("clause_id", ""))
        title = resolve_variables(title, variables)

        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.space_after = Pt(6)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

    # ── Clause Body ─────────────────────────────────────────────────────────

    def _add_clause(
        self, doc: Document, clause: dict, variables: dict, clause_number: int
    ) -> None:
        """
        Add a numbered clause to the document. Uses customText if present,
        otherwise resolves templateText.
        """
        # Determine the text to use
        raw_text = (
            clause.get("customText")
            or clause.get("custom_text")
            or clause.get("templateText")
            or clause.get("template_text")
            or ""
        )
        resolved = resolve_variables(raw_text, variables)

        # Clause number + title
        title = clause.get("title", "")
        if title:
            title = resolve_variables(title, variables)
            heading_p = doc.add_paragraph()
            heading_p.space_before = Pt(12)
            heading_p.space_after = Pt(4)
            run = heading_p.add_run(f"{clause_number}.  {title}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

        # Body text — render structured blocks so (a)/(i) markers, nested
        # numbering, and hanging indents survive to Word. A flat run dump (the
        # old html_to_docx_runs) collapsed the whole clause into one paragraph.
        blocks = clause_body_to_blocks(resolved)
        if not blocks:  # parse failure or empty — plain fallback
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.space_after = Pt(6)
            if not title:
                nr = p.add_run(f"{clause_number}.  ")
                nr.bold = True
                nr.font.size = Pt(12)
                nr.font.name = "Times New Roman"
            r = p.add_run(_strip_html(resolved))
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"
            return

        for idx, block in enumerate(blocks):
            p = doc.add_paragraph()
            p.space_after = Pt(6)
            level = block.get("indent", 1)
            p.paragraph_format.left_indent = Inches(0.5 * level)
            marker = block.get("marker")
            if marker:
                # Hanging indent: marker sits 0.5in left of the wrapped body.
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5 * level))
            align = block.get("align")
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Untitled clauses carry their number on the first line.
            if idx == 0 and not title:
                nr = p.add_run(f"{clause_number}.  ")
                nr.bold = True
                nr.font.size = Pt(12)
                nr.font.name = "Times New Roman"
            if marker:
                mr = p.add_run(f"{marker}\t")
                mr.font.size = Pt(12)
                mr.font.name = "Times New Roman"
            for text, bold, italic, underline in block["runs"]:
                run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
                if bold or block.get("heading"):
                    run.bold = True
                if italic:
                    run.italic = True
                if underline:
                    run.underline = True

    # ── Signing Pages ───────────────────────────────────────────────────────

    def _create_signing_page(
        self,
        doc: Document,
        document_type: str,
        signing_data: dict,
        variables: dict,
        page_break: bool = True,
    ) -> None:
        """
        Create the signing page appropriate for the document type.
        Uses Word tables for precise spacing (critical for legal documents).
        """
        if page_break:
            doc.add_page_break()

        if document_type in WILL_TYPES:
            self._signing_page_will(doc, document_type, signing_data, variables)
        elif document_type in POA_TYPES:
            self._signing_page_poa(doc, document_type, signing_data, variables)
        elif document_type in AFFIDAVIT_TYPES:
            self._signing_page_affidavit(doc, document_type, signing_data, variables)

    def _signing_page_will(
        self,
        doc: Document,
        document_type: str,
        signing_data: dict,
        variables: dict,
    ) -> None:
        """Testimonium block for wills with witness attestation."""
        testator_name = (
            variables.get("testatorFullName")
            or variables.get("testator_full_name")
            or "[TESTATOR NAME]"
        )
        pronoun = variables.get("pronoun", "his/her")
        if document_type == "probate_will":
            will_type_label = "Primary"
        elif document_type == "non_probate_will":
            will_type_label = "Non-Probate"
        else:
            will_type_label = ""  # single/short wills need no qualifier
        num_pages = signing_data.get("numberOfPages", variables.get("numberOfPages", "[__]"))

        # Testimonium heading
        p = doc.add_paragraph()
        p.space_before = Pt(24)
        run = p.add_run("TESTIMONIUM")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

        # Attestation table: left column = text with ), right column = signature
        table = doc.add_table(rows=7, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(3.0)

        # Remove table borders
        self._remove_table_borders(table)

        # Row 0
        left_lines = [
            f"SIGNED, PUBLISHED AND DECLARED  )",
            f"by the said Testator, {testator_name.upper()}  )",
            f"as {pronoun} {(will_type_label + ' ') if will_type_label else ''}  )",
            f"Will, consisting of {num_pages} pages  )",
            f"including this page, in the  )",
            f"presence of us, both present  )",
            f"at the same time, who at  )",
            f"{pronoun} request, in {pronoun} presence  )",
            f"and in the presence of each  )",
            f"other have hereunto subscribed  )",
            f"our names as witnesses.  )",
        ]

        # Fill left column across rows
        for i, line in enumerate(left_lines):
            if i < len(table.rows):
                cell = table.cell(i, 0)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(line)
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        # If we need more rows, add them
        while len(left_lines) > len(table.rows):
            table.add_row()

        # Re-fill if rows were added
        if len(left_lines) > 7:
            for i in range(7, len(left_lines)):
                cell = table.cell(i, 0)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(left_lines[i])
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        # Right column: signature line in the middle row area
        sig_row = 4
        cell = table.cell(sig_row, 1)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 30)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        cell = table.cell(sig_row + 1, 1)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(testator_name.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        doc.add_paragraph()  # spacer

        # Witness block
        p = doc.add_paragraph()
        run = p.add_run("WITNESSES:")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        # Two witness blocks side by side
        w_table = doc.add_table(rows=4, cols=2)
        w_table.columns[0].width = Inches(3.25)
        w_table.columns[1].width = Inches(3.25)
        self._remove_table_borders(w_table)

        for col in range(2):
            labels = ["Name:", "Occupation:", "Address:", ""]
            for row_idx, label in enumerate(labels):
                cell = w_table.cell(row_idx, col)
                cell.text = ""
                p = cell.paragraphs[0]
                if label:
                    run = p.add_run(f"{label}  {'_' * 25}")
                    run.font.size = Pt(11)
                    run.font.name = "Times New Roman"

    def _signing_page_poa(
        self,
        doc: Document,
        document_type: str,
        signing_data: dict,
        variables: dict,
    ) -> None:
        """Signature block for Powers of Attorney."""
        grantor_name = (
            variables.get("testatorFullName")
            or variables.get("testator_full_name")
            or "[GRANTOR NAME]"
        )
        doc_label = (
            "Continuing Power of Attorney for Property"
            if document_type == "poa_property"
            else "Power of Attorney for Personal Care"
        )

        # Heading
        p = doc.add_paragraph()
        p.space_before = Pt(24)
        run = p.add_run("EXECUTION")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

        # Execution text
        p = doc.add_paragraph()
        p.space_before = Pt(12)
        text = (
            f"IN WITNESS WHEREOF, I, {grantor_name.upper()}, have hereunto set "
            f"my hand and seal to this {doc_label}."
        )
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        # Signature table
        table = doc.add_table(rows=5, cols=2)
        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(3.0)
        self._remove_table_borders(table)

        # Left: date line
        cell = table.cell(0, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(f"Date: {'_' * 25}")
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        # Right: signature line
        cell = table.cell(0, 1)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 30)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        cell = table.cell(1, 1)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(grantor_name.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        # Witness block
        p = doc.add_paragraph()
        run = p.add_run("WITNESS:")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        for label in ["Name:", "Occupation:", "Address:"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}  {'_' * 40}")
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

    def _signing_page_affidavit(
        self,
        doc: Document,
        document_type: str,
        signing_data: dict,
        variables: dict,
    ) -> None:
        """Jurat and commissioner block for affidavits of execution."""
        deponent_name = (
            signing_data.get("deponentName")
            or variables.get("deponentName")
            or variables.get("otherWitnessName")
            or variables.get("other_witness_name")
            or "[DEPONENT NAME]"
        )
        commissioner_name = (
            signing_data.get("commissionerName")
            or variables.get("commissionerName")
            or variables.get("commissioner_name")
            or "[COMMISSIONER NAME]"
        )
        city = (
            variables.get("city")
            or variables.get("cityName")
            or variables.get("city_name")
            or "[City]"
        )
        province = variables.get("province", "Ontario")

        # Heading
        p = doc.add_paragraph()
        p.space_before = Pt(24)
        run = p.add_run("AFFIDAVIT OF EXECUTION")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Affidavit body
        p = doc.add_paragraph()
        run = p.add_run(
            f"I, {deponent_name.upper()}, of the City of {city}, "
            f"in the Province of {province}, MAKE OATH AND SAY:"
        )
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        # Numbered affidavit paragraphs
        testator_name = (
            variables.get("testatorFullName")
            or variables.get("testator_full_name")
            or "[TESTATOR NAME]"
        )
        pronoun = variables.get("pronoun", "his/her")

        affidavit_paras = [
            (
                f"1.  I was present and saw {testator_name.upper()} named "
                f"in the within instrument sign the instrument."
            ),
            (
                f"2.  {testator_name.upper()} signed the instrument in my "
                f"presence and in the presence of the other subscribing witness."
            ),
            (
                f"3.  I know {testator_name.upper()} to be the person whose name "
                f"is signed to the instrument as its maker."
            ),
            f"4.  I am the subscribing witness to the instrument.",
            (
                f"5.  I am over the age of eighteen (18) years and I believe that "
                f"I am mentally competent to be a witness."
            ),
        ]

        for text in affidavit_paras:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.space_after = Pt(6)
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

        doc.add_paragraph()

        # Jurat table
        jurat_table = doc.add_table(rows=6, cols=2)
        jurat_table.columns[0].width = Inches(3.5)
        jurat_table.columns[1].width = Inches(3.0)
        self._remove_table_borders(jurat_table)

        # Left: jurat text
        jurat_lines = [
            f"SWORN BEFORE ME at the",
            f"City of {city},",
            f"in the Province of {province},",
            f"this _____ day of __________, 20____",
            "",
            "",
        ]
        for i, line in enumerate(jurat_lines):
            cell = jurat_table.cell(i, 0)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

        # Right: deponent signature
        cell = jurat_table.cell(2, 1)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 30)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        cell = jurat_table.cell(3, 1)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(deponent_name.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        # Commissioner line at bottom left
        cell = jurat_table.cell(4, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run("_" * 30)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        cell = jurat_table.cell(5, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(f"A Commissioner for Taking Affidavits")
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        # Commissioner name
        p = doc.add_paragraph()
        p.space_before = Pt(6)
        run = p.add_run(f"Name: {commissioner_name}")
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        # Exhibit A cover page for affidavits
        self._create_exhibit_a_cover(doc, document_type, variables)

    def _create_exhibit_a_cover(
        self, doc: Document, document_type: str, variables: dict
    ) -> None:
        """Exhibit 'A' cover page for affidavits."""
        doc.add_page_break()

        deponent_name = (
            variables.get("deponentName")
            or variables.get("otherWitnessName")
            or variables.get("other_witness_name")
            or "[DEPONENT NAME]"
        )
        commissioner_name = (
            variables.get("commissionerName")
            or variables.get("commissioner_name")
            or "[COMMISSIONER NAME]"
        )

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(72)
        run = p.add_run('THIS IS EXHIBIT "A"')
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"referred to in the Affidavit of {deponent_name.upper()}"
        )
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"sworn before me this _____ day of __________, 20____")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 35)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("A Commissioner for Taking Affidavits")
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # ── Schedule A ──────────────────────────────────────────────────────────

    def _create_schedule_a(self, doc: Document, variables: dict) -> None:
        """
        Schedule 'A' for probate will: lists Excluded Property definition
        and cross-references the Non-Probate Will.
        """
        doc.add_page_break()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(24)
        run = p.add_run('SCHEDULE "A"')
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("EXCLUDED PROPERTY")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        testator_name = (
            variables.get("testatorFullName")
            or variables.get("testator_full_name")
            or "[TESTATOR NAME]"
        )

        p = doc.add_paragraph()
        text = (
            f'"Excluded Property" means all property owned by '
            f"{testator_name.upper()} at the date of death that does not "
            f"require a Certificate of Appointment of Estate Trustee "
            f"(probate) for its transfer or realization, including but "
            f"not limited to:"
        )
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        items = [
            "Property held in joint tenancy with right of survivorship;",
            "Property with designated beneficiaries (life insurance, RRSPs, RRIFs, TFSAs, pensions);",
            "Property held in trust;",
            "Personal effects and household contents;",
            "Any other property that can be transferred or realized without probate.",
        ]

        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.space_after = Pt(4)
            run = p.add_run(f"\u2022  {item}")
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

        doc.add_paragraph()

        p = doc.add_paragraph()
        text = (
            "The Excluded Property is dealt with under the Testator's "
            "Non-Probate Will of even date."
        )
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    # ── Page Numbers ────────────────────────────────────────────────────────

    def _add_page_numbers(self, doc: Document) -> None:
        """Add 'Page X of Y' footer to all sections (right-aligned)."""
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Page number field
            run = p.add_run()
            fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._r.append(fld_char_begin)

            run2 = p.add_run()
            instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run2._r.append(instr)

            run3 = p.add_run()
            fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
            run3._r.append(fld_char_sep)

            run4 = p.add_run()
            fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run4._r.append(fld_char_end)

            run5 = p.add_run(" of ")
            run5.font.size = Pt(10)
            run5.font.name = "Times New Roman"

            # Total pages field
            run6 = p.add_run()
            fld2_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run6._r.append(fld2_begin)

            run7 = p.add_run()
            instr2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
            run7._r.append(instr2)

            run8 = p.add_run()
            fld2_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
            run8._r.append(fld2_sep)

            run9 = p.add_run()
            fld2_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run9._r.append(fld2_end)

    # ── Document Styles ─────────────────────────────────────────────────────

    def _set_document_styles(self, doc: Document, compact: bool = False) -> None:
        """Set base document styles: font, margins, line spacing."""
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11 if compact else 12)
        style.paragraph_format.line_spacing = 1.15 if compact else 1.5

        margin = Inches(0.75 if compact else 1)
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # Heading styles
        for level in range(1, 4):
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
                h_style.font.name = "Times New Roman"
                h_style.font.size = Pt(14)
                h_style.font.bold = True

    # ── Table Helpers ───────────────────────────────────────────────────────

    @staticmethod
    def _remove_table_borders(table) -> None:
        """Remove all borders from a table for clean signing-page layout."""
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}/>'
        )
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        tbl_pr.append(borders)

    # ── Batch Generation ────────────────────────────────────────────────────

    def generate_all_documents(
        self,
        draft_data: dict,
        clause_selections: dict[str, list[dict]],
        variables: dict,
    ) -> dict[str, bytes]:
        """
        Generate all required documents for a client.

        Args:
            draft_data: The full draft record from the database.
            clause_selections: Dict mapping document_type -> list of clause dicts.
            variables: Resolved placeholder variables.

        Returns:
            Dict mapping document_type -> DOCX file bytes.
        """
        results: dict[str, bytes] = {}

        for doc_type, clauses in clause_selections.items():
            if not clauses:
                continue
            try:
                docx_bytes = self.generate_document(
                    document_type=doc_type,
                    clauses=clauses,
                    variables=variables,
                )
                results[doc_type] = docx_bytes
                logger.info("Generated %s (%d bytes)", doc_type, len(docx_bytes))
            except Exception:
                logger.exception("Failed to generate %s", doc_type)

        return results
